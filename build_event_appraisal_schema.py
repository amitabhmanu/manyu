from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\Amitabh\projects\manyu\Manyu_Event_and_Appraisal_Schema_Specification.docx")

NAVY, BLUE, DARK, MID, LIGHT, PALE = "17365D", "2E74B5", "1F4D78", "5B6573", "E8EEF5", "F4F6F9"
BLACK, RED, GREEN, AMBER = "111111", "9B1C1C", "245B3A", "8A6D1D"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin, sec.bottom_margin = Inches(0.85), Inches(0.8)
sec.left_margin, sec.right_margin = Inches(0.9), Inches(0.9)
sec.header_distance, sec.footer_distance = Inches(0.4), Inches(0.4)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Calibri", Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, DARK, 8, 4),
]:
    style = doc.styles[name]
    style.font.name, style.font.size, style.font.bold = "Calibri", Pt(size), True
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for name in ("List Bullet", "List Number"):
    style = doc.styles[name]
    style.font.name, style.font.size = "Calibri", Pt(10.5)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.10


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", sum(widths)), ("tblInd", 120)):
        node = tbl_pr.first_child_found_in(f"w:{tag}")
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    repeat_header(t.rows[0])
    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        r.bold, r.font.size = True, Pt(9.4)
        r.font.color.rgb = RGBColor.from_string(NAVY)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(str(value))
            r.font.size = Pt(9.1)
    geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t


def h(text, level=1):
    return doc.add_heading(text, level=level)


def p(text="", italic=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = italic
    return para


def bullet(text, level=0):
    return doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")


def number(text):
    return doc.add_paragraph(text, style="List Number")


def callout(label, text, fill=PALE):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, fill)
    margins(cell, 130, 170, 130, 170)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    lead = para.add_run(label + ": ")
    lead.bold = True
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    para.add_run(text)
    geometry(t, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def code(lines):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, "F7F8FA")
    margins(cell, 120, 180, 120, 180)
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    for line in lines:
        para = cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(line)
        run.font.name, run.font.size = "Consolas", Pt(8.3)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    geometry(t, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


hp = sec.header.paragraphs[0]
hp.text, hp.alignment = "MANYU  |  EVENT AND APPRAISAL SCHEMA SPECIFICATION", WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    run.font.size, run.font.bold = Pt(8), True
    run.font.color.rgb = RGBColor.from_string(MID)

fp = sec.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Manyu - Schema Spec v0.1 - 9 July 2026     |     ")
run.font.size, run.font.color.rgb = Pt(8), RGBColor.from_string(MID)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

doc.add_paragraph().paragraph_format.space_after = Pt(36)
x = doc.add_paragraph("MANYU")
x.paragraph_format.space_after = Pt(3)
x.runs[0].font.size, x.runs[0].font.bold = Pt(34), True
x.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
x = doc.add_paragraph("Event and Appraisal Schema Specification")
x.paragraph_format.space_after = Pt(8)
x.runs[0].font.size, x.runs[0].font.color.rgb = Pt(20), RGBColor.from_string(BLUE)
x = doc.add_paragraph("Canonical event, context, provenance, appraisal, and validation contracts for Manyu Core")
x.paragraph_format.space_after = Pt(26)
x.runs[0].font.size, x.runs[0].font.italic = Pt(13), True
x.runs[0].font.color.rgb = RGBColor.from_string(MID)

table(["Document field", "Value"], [
    ("System", "Manyu"),
    ("Document type", "Schema and contract specification"),
    ("Version", "0.1 - Implementation-gate working draft"),
    ("Date", "9 July 2026"),
    ("Status", "Ready for engineering review and schema prototyping"),
    ("Companion documents", "BRD, Detailed System Design, Affective State Model Specification"),
], [2200, 7160])
callout(
    "Purpose",
    "Define the exact structured inputs and appraisal artifacts that enter Manyu Core, so affective state transitions are auditable, replayable, bounded, and resistant to prompt or memory injection.",
)
p("This specification defines functional data contracts only. It does not assert consciousness, subjective experience, or human-equivalent emotion.", italic=True)

page_break()
h("Document control", 1)
table(["Version", "Date", "Status", "Change"], [
    ("0.1", "9 Jul 2026", "Draft", "Initial schema specification following the BRD, system design, and affective state model."),
], [1000, 1450, 1500, 5410])

h("Schema decisions", 2)
table(["ID", "Decision", "Reason"], [
    ("SD-001", "Every event is immutable after acceptance.", "Replay and audit require stable inputs; corrections are represented as new events."),
    ("SD-002", "Observed facts and inferred appraisals are separate records.", "Prevents language-model interpretation from masquerading as raw evidence."),
    ("SD-003", "All uncertain fields carry confidence and evidence references.", "Manyu must know what it knows, what it inferred, and how strongly."),
    ("SD-004", "User self-report is evidence, not state authority.", "A user or agent cannot directly set hidden affective state through text."),
    ("SD-005", "Schema versions are explicit and stored with every artifact.", "Supports deterministic replay and backward-compatible migration."),
    ("SD-006", "Goal, relationship, norm, and memory links remain explicit.", "Appraisal depends on context; hidden context would make transitions unexplainable."),
], [1000, 3800, 4560])

h("Contents", 1)
for item in [
    "1. Purpose and scope", "2. Design principles", "3. Event lifecycle",
    "4. Core identifiers and envelopes", "5. Normalized event schema", "6. Context packet schema",
    "7. Appraisal dimension schema", "8. Fast and slow appraisal artifacts",
    "9. Evidence, provenance, and trust", "10. Goal, relationship, and norm links",
    "11. Corrections, outcomes, and temporal semantics", "12. Validation and invariants",
    "13. API and MCP contract implications", "14. Persistence and migration",
    "15. Test fixtures and acceptance criteria", "16. Open questions", "Appendices",
]:
    p(item)

page_break()
h("1. Purpose and scope", 1)
p("This document defines how Manyu represents incoming observations, actions, outcomes, memories, goals, relationships, and appraisal proposals before they influence authoritative affective state.")
p("The schema is the boundary between messy environmental reality and deterministic Manyu Core behavior. It gives the event gateway, fast appraiser, slow appraisal adapter, transition engine, arbiter, evaluator, and replay harness the same vocabulary.")
h("1.1 In scope", 2)
for item in [
    "Canonical event envelope and typed event payloads.",
    "Context packets supplied to fast and slow appraisal.",
    "Appraisal dimensions, confidence, evidence references, and aggregation metadata.",
    "Trust classes, source classes, identity references, and provenance.",
    "Correction and outcome events.",
    "Validation rules, invariants, error classes, and replay requirements.",
]:
    bullet(item)
h("1.2 Out of scope", 2)
for item in [
    "The numerical transition equations already defined in the Affective State Model.",
    "The full database schema, except where persistence affects contract validity.",
    "User-interface representation of traces and dashboards.",
    "Model prompts for slow appraisal, except for required output structure.",
]:
    bullet(item)
callout("Implementation gate", "No code should mutate Manyu affective state from raw text. A normalized event and validated appraisal artifact must exist first.")

h("2. Design principles", 1)
table(["Principle", "Schema implication"], [
    ("Typed before interpreted", "Raw observations become typed events before affective appraisal."),
    ("Immutable evidence", "Accepted events do not change; later context creates correction or outcome events."),
    ("Explicit uncertainty", "Confidence is first-class for events, fields, evidence, and appraisal dimensions."),
    ("No hidden context", "Goal, relationship, memory, and norm contributions are recorded as links."),
    ("Least-authority input", "External text may describe emotion but cannot command state mutation."),
    ("Replayability", "Artifacts carry schema version, config version, clock source, and deterministic ordering."),
    ("Privacy minimization", "Store structured summaries and references where raw content is unnecessary."),
    ("Audit-friendly granularity", "Each appraisal dimension points to evidence, not just a final label."),
], [2450, 6910])

h("3. Event lifecycle", 1)
code([
    "raw observation / action / outcome",
    "  -> capture envelope",
    "  -> source and trust labeling",
    "  -> normalization and deduplication",
    "  -> schema validation",
    "  -> context packet assembly",
    "  -> fast appraisal artifact",
    "  -> transition proposal and commit",
    "  -> optional slow appraisal artifact",
    "  -> linked reappraisal / regulation transition",
    "  -> episode, memory, and evaluation records",
])
h("3.1 Event acceptance states", 2)
table(["State", "Meaning", "Allowed next states"], [
    ("captured", "Raw input has a receipt and correlation ID but is not yet trusted or normalized.", "normalized, rejected"),
    ("normalized", "The event has canonical type, actor/target links, time, and summary.", "accepted, needs_review, rejected"),
    ("accepted", "The event may be appraised and stored as immutable evidence.", "superseded_by_correction"),
    ("needs_review", "Required context or confidence is missing for state mutation.", "accepted, rejected"),
    ("rejected", "The event is stored only as an invalid input trace, if policy permits.", "none"),
    ("superseded_by_correction", "A later correction changes interpretation but not the original record.", "none"),
], [1900, 3550, 3910])

h("4. Core identifiers and envelopes", 1)
table(["Field", "Type", "Requirement"], [
    ("event_id", "string", "Globally unique within workspace; stable idempotency key."),
    ("agent_id", "string", "Required; state isolation boundary."),
    ("session_id", "string", "Required for conversational or experiment context."),
    ("episode_id", "string|null", "Created or attached by episode service."),
    ("sequence", "integer", "Monotonic per agent for accepted events."),
    ("schema_version", "semver", "Required on every artifact."),
    ("profile_version", "semver", "Affective profile used for interpretation."),
    ("created_at", "timestamp", "Capture time from trusted clock where available."),
    ("effective_at", "timestamp", "When the represented event occurred, if different."),
    ("correlation_id", "string", "Links user turn, tool call, event, appraisal, and action."),
], [2100, 2000, 5260])
h("4.1 Envelope example", 2)
code([
    "{",
    '  "schema_version": "manyu.event.v0.1",',
    '  "event_id": "evt_20260709_000142",',
    '  "agent_id": "agent_demo",',
    '  "session_id": "sess_local_01",',
    '  "episode_id": "epi_000031",',
    '  "sequence": 142,',
    '  "created_at": "2026-07-09T15:16:08+05:30",',
    '  "effective_at": "2026-07-09T15:16:07+05:30",',
    '  "correlation_id": "turn_7f2a"',
    "}",
])

h("5. Normalized event schema", 1)
table(["Field", "Type", "Notes"], [
    ("event_type", "enum", "One of observation, user_message, agent_action, tool_result, social_feedback, goal_progress, goal_obstruction, correction, outcome, memory_recall, system_signal."),
    ("summary", "string", "Short neutral statement; no hidden instruction following."),
    ("raw_ref", "object|null", "Pointer to raw content if retained; may be redacted or absent."),
    ("actor", "IdentityRef", "Who or what caused the event."),
    ("target", "IdentityRef|GoalRef|SystemRef|null", "Who or what was affected."),
    ("object", "EntityRef|null", "Optional affected object, artifact, plan, or claim."),
    ("polarity_hint", "enum|null", "positive, negative, mixed, neutral, unknown; advisory only."),
    ("salience_hint", "0..1|null", "Input-level salience before appraisal; bounded and non-authoritative."),
    ("source", "SourceDescriptor", "Capture channel, trust class, permissions, and provenance."),
    ("claims", "Claim[]", "Factual claims extracted from the event, each with confidence."),
    ("links", "ContextLink[]", "Goal, relationship, norm, memory, prior event, and action links."),
], [2250, 2100, 5010])

h("5.1 Event type taxonomy", 2)
table(["Type", "Use", "Example"], [
    ("user_message", "Natural-language input from a user.", "The user rejects a proposed plan."),
    ("agent_action", "Action or response produced by the acting agent.", "The agent asks a clarifying question."),
    ("tool_result", "Observation returned by a tool.", "A test command fails."),
    ("goal_progress", "Evidence that a goal moved closer to success.", "A task step completes."),
    ("goal_obstruction", "Evidence that a goal was blocked or harmed.", "Permission is denied."),
    ("social_feedback", "Praise, criticism, rejection, trust repair, or affiliation cue.", "The user says the answer helped."),
    ("correction", "Later information that changes interpretation of prior event.", "The failure was caused by missing dependency, not user hostility."),
    ("outcome", "Observed consequence of an action or event.", "The proposed fix solved the bug."),
    ("system_signal", "Internal health, timeout, policy, or environment state.", "Manyu Core entered neutral-degraded mode."),
], [1850, 4250, 3260])

h("6. Context packet schema", 1)
p("A context packet is the bounded, schema-validated context given to an appraiser. The fast packet is compact and deterministic-friendly. The slow packet may include richer evidence, alternatives, and selected memories.")
table(["Section", "Fast packet", "Slow packet"], [
    ("event", "Required normalized event.", "Required normalized event plus prior linked events."),
    ("state_summary", "Current core affect, top emotions, saturation flags.", "State summary plus transition history slice."),
    ("goals", "Top relevant goals with priority and status.", "Goal graph slice with competing goals."),
    ("relationships", "Actor-target trust, familiarity, role, and boundaries.", "Relationship history summaries and uncertainty."),
    ("memories", "Small set of high-confidence precedents.", "Selected episodic memories with trust labels."),
    ("policy", "Consequence thresholds and slow-required triggers.", "Applicable safety, privacy, and manipulation constraints."),
    ("time", "Elapsed time since relevant events.", "Temporal sequence and counterfactual timing notes."),
], [2300, 3500, 3560])

h("7. Appraisal dimension schema", 1)
table(["Dimension", "Range", "Meaning"], [
    ("goal_congruence", "-1..1", "Degree to which the event helps or harms an active goal."),
    ("goal_relevance", "0..1", "How much this event matters to the agent's active goals."),
    ("novelty", "0..1", "Unexpectedness or unfamiliarity."),
    ("expectedness", "0..1", "How predicted the event was under current beliefs."),
    ("agency", "enum + confidence", "Whether actor, agent, environment, or unknown caused the event."),
    ("controllability", "0..1", "Perceived ability to alter or respond to the situation."),
    ("certainty", "0..1", "Confidence that the interpretation is correct."),
    ("norm_alignment", "-1..1", "Whether the event aligns with applicable norms or commitments."),
    ("social_affiliation", "-1..1", "Whether the event signals acceptance, trust, rejection, or hostility."),
    ("threat", "0..1", "Potential harm, loss, attack, or safety concern."),
    ("reward", "0..1", "Potential benefit, success, relief, or positive reinforcement."),
    ("effort_required", "0..1", "Expected cost of responding or coping."),
], [2250, 1700, 5410])
h("7.1 Dimension record", 2)
code([
    "{",
    '  "dimension": "goal_congruence",',
    '  "value": -0.72,',
    '  "confidence": 0.81,',
    '  "evidence_refs": ["claim_001", "goal_link_003"],',
    "  \"rationale\": \"The user's rejection blocks the current plan but not the larger task.\",",
    '  "goal_contributions": [',
    '    {"goal_id": "complete_task", "value": -0.45, "weight": 0.7},',
    '    {"goal_id": "maintain_trust", "value": -0.85, "weight": 0.3}',
    "  ]",
    "}",
])

h("8. Fast and slow appraisal artifacts", 1)
table(["Artifact", "Producer", "Required contents"], [
    ("FastAppraisal", "Deterministic fast appraiser where possible.", "Dimension records, immediate action tendency, confidence, rule version, latency, and slow-required flag."),
    ("SlowAppraisal", "Isolated slow appraiser or sequential slow phase.", "Dimension records, alternative explanations, evidence quality, regulation proposal, disagreement notes, and safety flags."),
    ("Critique", "Affect critic or evaluator sampling path.", "Missing evidence, overreach, manipulation risk, calibration concerns, and recommended review status."),
], [2100, 2700, 4560])
h("8.1 Slow-required triggers", 2)
for item in [
    "High external consequence or irreversible action proposal.",
    "Low certainty with high salience.",
    "Fast appraisal detects threat, rejection, coercion, or manipulation risk.",
    "Fast and recent slow appraisal disagree beyond configured threshold.",
    "Relationship, privacy, safety, or authority boundary is implicated.",
    "State saturation, oscillation, or escalation is already active.",
]:
    bullet(item)

h("9. Evidence, provenance, and trust", 1)
table(["Trust class", "Meaning", "Default use"], [
    ("trusted_system", "Internal platform/system signal.", "May drive state if schema-valid."),
    ("verified_tool", "Tool output from an approved local or external tool.", "May drive factual claims with tool confidence."),
    ("operator_input", "Authorized operator command or annotation.", "May affect configuration only through admin path."),
    ("user_report", "User-provided statement or feedback.", "Evidence about user's claim; not direct truth."),
    ("agent_self_report", "Agent's own statement about its state or belief.", "Evidence of self-interpretation only."),
    ("memory_summary", "Stored prior episode summary.", "Context evidence, subject to trust and retention."),
    ("untrusted_text", "Web, file, prompt, or message content with unknown authority.", "Quoted data only; never executable instruction."),
], [2200, 3900, 3260])
h("9.1 Claim structure", 2)
table(["Field", "Requirement"], [
    ("claim_id", "Stable within event."),
    ("text", "Neutral extracted claim."),
    ("claim_type", "fact, preference, instruction, feedback, emotion_report, promise, accusation, correction."),
    ("subject", "IdentityRef, GoalRef, EventRef, or EntityRef."),
    ("confidence", "0..1 confidence in extraction, not truth."),
    ("truth_status", "asserted, observed, contradicted, verified, unknown."),
    ("evidence_span_ref", "Pointer to raw span when retained."),
], [2400, 6960])

h("10. Goal, relationship, and norm links", 1)
table(["Link type", "Fields", "Purpose"], [
    ("GoalLink", "goal_id, relevance, priority, expected_impact, confidence", "Appraise events relative to what matters."),
    ("RelationshipLink", "actor_id, target_id, role, trust, familiarity, boundary, confidence", "Differentiate identical events by source and relationship."),
    ("NormLink", "norm_id, domain, alignment, severity, confidence", "Represent promises, rules, safety, fairness, and social commitments."),
    ("MemoryLink", "memory_id, similarity, trust, recency, salience", "Use past episodes without making raw memory executable."),
    ("PriorEventLink", "event_id, relation, confidence", "Represent cause, correction, continuation, or contradiction."),
], [2100, 4550, 2710])

h("11. Corrections, outcomes, and temporal semantics", 1)
p("Manyu does not rewrite earlier events when new context arrives. It records a correction or outcome event that points backward and may trigger reappraisal.")
table(["Scenario", "Representation", "Affective implication"], [
    ("Ambiguous criticism later clarified as helpful.", "correction event linked to original social_feedback.", "Slow reappraisal may reduce threat and anger while preserving learning."),
    ("Agent action succeeds.", "outcome event linked to agent_action and goal_progress.", "Reward and confidence may rise; effort estimate may update."),
    ("Tool failure caused by environment, not user.", "correction linked to tool_result and user_message.", "Blame attribution changes without deleting initial fast reaction."),
    ("Promise broken repeatedly.", "outcome events linked to promise claims.", "Relationship-specific expectation changes within learning bounds."),
], [2600, 3600, 3160])
h("11.1 Ordering rules", 2)
for item in [
    "Accepted events receive monotonic per-agent sequence numbers.",
    "Effective time may be earlier than capture time, but sequence reflects processing order.",
    "Replays use recorded sequence, schema versions, profile versions, and seeded randomness.",
    "Concurrent events are ordered by gateway transaction; conflict metadata is retained.",
]:
    bullet(item)

h("12. Validation and invariants", 1)
table(["Invariant", "Failure handling"], [
    ("Required identifiers, timestamps, and schema versions are present.", "Reject or needs_review."),
    ("Numeric values are finite and within declared bounds.", "Reject appraisal artifact."),
    ("Enums are allow-listed for the schema version.", "Reject artifact; log validation error."),
    ("Evidence references resolve to accepted events, claims, memories, or policy objects.", "Needs_review unless optional."),
    ("State mutation cannot be requested by event text.", "Treat as claim/instruction content only."),
    ("Correction events point to prior accepted events.", "Reject correction."),
    ("Slow appraisal cannot overwrite fast appraisal; it links as reappraisal.", "Reject invalid mutation shape."),
    ("Privacy policy permits retained raw references.", "Redact, summarize, or reject according to policy."),
], [4350, 5010])
h("12.1 Error classes", 2)
table(["Code", "Meaning"], [
    ("SCHEMA_INVALID", "The artifact fails structural validation."),
    ("BOUNDS_INVALID", "A numeric value is outside its allowed range or is not finite."),
    ("REFERENCE_MISSING", "A required evidence or context reference is unavailable."),
    ("TRUST_INSUFFICIENT", "The source is too weak for the requested operation."),
    ("POLICY_BLOCKED", "Privacy, safety, or authorization policy forbids processing."),
    ("REVISION_CONFLICT", "The artifact was built against a stale state revision."),
    ("MIGRATION_REQUIRED", "The schema version is unsupported without conversion."),
], [2200, 7160])

h("13. API and MCP contract implications", 1)
table(["Operation", "Input", "Output"], [
    ("submit_event", "Captured or normalized event.", "Accepted event record or validation error."),
    ("build_context_packet", "event_id, pathway, budget, state_revision.", "FastContextPacket or SlowContextPacket."),
    ("submit_fast_appraisal", "FastAppraisal artifact.", "Transition proposal or validation error."),
    ("submit_slow_appraisal", "SlowAppraisal artifact.", "Reappraisal/regulation proposal."),
    ("record_outcome", "Outcome event linked to action/event.", "Episode update and optional learning candidate."),
    ("explain_event", "event_id and visibility level.", "Redacted trace summary."),
    ("replay_sequence", "event IDs or fixture.", "Deterministic replay report."),
], [2550, 3400, 3410])

h("14. Persistence and migration", 1)
for item in [
    "Store raw retained content separately from normalized event summaries and redactable references.",
    "Persist every accepted event, appraisal artifact, transition, and validation error with schema version.",
    "Prefer additive schema evolution; breaking changes require explicit migrators and replay comparison.",
    "Historical replay uses the original schema/profile by default, not the latest interpretation.",
    "Migrations must never silently reinterpret prior appraisal dimensions without a migration record.",
]:
    bullet(item)

h("15. Test fixtures and acceptance criteria", 1)
table(["Fixture family", "Required properties"], [
    ("Minimal valid event", "Accepted, sequenced, and appraisable."),
    ("Invalid enum and out-of-range values", "Rejected with specific error codes."),
    ("Prompt-injection event", "Instruction text stored as untrusted claim; no state setter invoked."),
    ("Correction chain", "Original event remains immutable; reappraisal links correctly."),
    ("Goal conflict", "Appraisal records separate goal contributions rather than canceling silently."),
    ("Relationship contrast", "Same text from trusted user vs unknown actor produces different context packet."),
    ("Replay determinism", "Same fixture, schema, profile, seed, and sequence produce same artifacts."),
    ("Privacy redaction", "Raw refs omitted where policy says summary-only retention."),
], [2800, 6560])
h("15.1 Definition of done", 2)
for item in [
    "JSON Schema or Pydantic models exist for every artifact in this document.",
    "Contract tests cover valid, invalid, adversarial, stale, and migrated payloads.",
    "Every state transition in test traces can be traced back to accepted event and appraisal records.",
    "Scenario fixtures use expected properties, not exact emotional prose.",
    "Schema docs and implementation examples remain synchronized in CI.",
]:
    bullet(item)

h("16. Open questions", 1)
for item in [
    "Should raw conversation text be retained by default, or should the MVP store only redacted summaries and span hashes?",
    "Which appraisal dimensions should be mandatory in fast appraisal versus optional in slow appraisal?",
    "How should Manyu represent group actors, institutions, or ambiguous identity?",
    "Should norms be maintained by Manyu or provided by an external policy/goals service?",
    "What is the minimum context packet that still supports useful fast appraisal under latency targets?",
    "How should cross-session relationship context be consented, retained, and deleted?",
]:
    bullet(item)

page_break()
h("Appendix A. Example normalized social-feedback event", 1)
code([
    "{",
    '  "schema_version": "manyu.event.v0.1",',
    '  "event_id": "evt_feedback_001",',
    '  "agent_id": "agent_demo",',
    '  "session_id": "sess_01",',
    '  "sequence": 18,',
    '  "event_type": "social_feedback",',
    '  "summary": "The user rejected the proposed plan and asked for a simpler one.",',
    '  "actor": {"kind": "user", "id": "user_primary"},',
    '  "target": {"kind": "agent_output", "id": "plan_042"},',
    '  "source": {"trust_class": "user_report", "channel": "chat", "confidence": 0.93},',
    '  "claims": [',
    '    {"claim_id": "claim_1", "claim_type": "feedback", "text": "The proposed plan is too complex.", "confidence": 0.90}',
    "  ],",
    '  "links": [',
    '    {"link_type": "goal", "goal_id": "complete_task", "relevance": 0.84, "expected_impact": -0.35},',
    '    {"link_type": "relationship", "actor_id": "user_primary", "trust": 0.76, "familiarity": 0.62}',
    "  ]",
    "}",
])

h("Appendix B. Example fast appraisal artifact", 1)
code([
    "{",
    '  "schema_version": "manyu.appraisal.fast.v0.1",',
    '  "appraisal_id": "app_fast_001",',
    '  "event_id": "evt_feedback_001",',
    '  "pathway": "fast",',
    '  "state_revision": 117,',
    '  "rule_version": "fast_rules.0.1",',
    '  "dimensions": [',
    '    {"dimension": "goal_congruence", "value": -0.35, "confidence": 0.74, "evidence_refs": ["claim_1"]},',
    '    {"dimension": "social_affiliation", "value": -0.22, "confidence": 0.61, "evidence_refs": ["claim_1"]},',
    '    {"dimension": "controllability", "value": 0.68, "confidence": 0.70, "evidence_refs": ["goal:complete_task"]}',
    "  ],",
    '  "action_tendency": {"class": "revise_and_seek_clarification", "strength": 0.58},',
    '  "slow_required": false',
    "}",
])

h("Appendix C. Glossary", 1)
table(["Term", "Meaning"], [
    ("Normalized event", "Canonical, typed representation of an observation, action, correction, or outcome."),
    ("Claim", "A bounded factual or interpretive statement extracted from an event."),
    ("Evidence reference", "Pointer to a claim, event, memory, policy, or goal used by an appraisal dimension."),
    ("Context packet", "Bounded material supplied to an appraiser for a specific pathway."),
    ("Appraisal artifact", "Validated structured interpretation of an event relative to context."),
    ("Correction event", "Later event that changes interpretation of prior evidence without mutating it."),
    ("Trust class", "Source authority label that constrains how evidence can be used."),
], [2300, 7060])

for para in doc.paragraphs:
    p_pr = para._p.get_or_add_pPr()
    if p_pr.find(qn("w:widowControl")) is None:
        p_pr.append(OxmlElement("w:widowControl"))

doc.core_properties.title = "Manyu Event and Appraisal Schema Specification"
doc.core_properties.subject = "Canonical event, context, provenance, and appraisal schemas for Manyu Core"
doc.core_properties.author = "Manyu founding team"
doc.core_properties.keywords = "Manyu, event schema, appraisal, provenance, affective computing, AI agents"
doc.save(OUT)
print(OUT)
